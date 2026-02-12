"""apache_age.txt를 서식이 적용된 Word 문서(.docx)로 변환하는 스크립트"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

INPUT = Path(__file__).parent / "apache_age.txt"
OUTPUT = Path(__file__).parent / "apache_age.docx"

FONT_KO = "맑은 고딕"
FONT_CODE = "Consolas"


def set_font_east_asia(run, font_name):
    """run의 동아시아(한글) 폰트를 설정한다."""
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name):
    """스타일의 서양 + 동아시아 폰트를 모두 설정한다."""
    style.font.name = font_name
    # 동아시아 폰트 설정
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, color_hex):
    """테이블 셀 배경색 설정"""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    tc_pr.append(shading_elem)


def make_table_from_lines(doc, lines):
    """텍스트 테이블 라인을 Word 테이블로 변환"""
    data_rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 구분선(───, ├──) 건너뛰기
        if re.match(r"^[┌├└─┬┼┴┐┤┘│\s]+$", line):
            continue
        # │ 로 분리
        cells = [c.strip() for c in re.split(r"[│]", line) if c.strip()]
        if cells:
            data_rows.append(cells)

    if not data_rows:
        return

    num_cols = max(len(r) for r in data_rows)
    # 모든 행의 열 수 맞추기
    for r in data_rows:
        while len(r) < num_cols:
            r.append("")

    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles["Normal"]
                for run in paragraph.runs:
                    run.font.name = FONT_KO
                    set_font_east_asia(run, FONT_KO)
                    run.font.size = Pt(9)
            if i == 0:
                # 헤더 행 스타일
                set_cell_shading(cell, "2B579A")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.name = FONT_KO
                        set_font_east_asia(run, FONT_KO)
                        run.font.size = Pt(9)

    return table


def is_table_line(line):
    """테이블의 일부인 라인인지 판별"""
    stripped = line.strip()
    if not stripped:
        return False
    return bool(re.match(r"^[┌├└│]", stripped))


def is_section_header(line):
    """=== 로 둘러싸인 섹션 제목 판별"""
    return line.strip().startswith("=") and len(line.strip()) > 10


def is_subsection_divider(line):
    """──── 으로 된 소제목 구분선 판별"""
    stripped = line.strip()
    return stripped.startswith("──") and len(stripped) > 10


def add_code_block(doc, code_lines):
    """코드 블록을 회색 배경의 단락으로 추가"""
    code_text = "\n".join(code_lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = FONT_CODE
    set_font_east_asia(run, FONT_CODE)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # 배경색 설정
    p_pr = p._element.get_or_add_pPr()
    shading_elem = p_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): "F0F0F0", qn("w:val"): "clear"},
    )
    p_pr.append(shading_elem)


def build_docx():
    text = INPUT.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    set_style_font(style, FONT_KO)
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    # 제목 스타일
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        set_style_font(hs, FONT_KO)
        hs.font.color.rgb = RGBColor(43, 87, 154)

    # 리스트 스타일
    for sn in ("List Bullet", "List Number"):
        if sn in doc.styles:
            set_style_font(doc.styles[sn], FONT_KO)

    # ── 문서 제목 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run("Apache AGE (A Graph Extension)")
    run.font.name = FONT_KO
    set_font_east_asia(run, FONT_KO)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(43, 87, 154)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("종합 자료")
    run.font.name = FONT_KO
    set_font_east_asia(run, FONT_KO)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    i = 0
    # 첫 번째 === 블록(타이틀)과 목차 이후부터 시작
    # 목차 끝 찾기
    while i < len(lines):
        if "1.  개요" in lines[i]:
            break
        i += 1

    # 목차 수집
    toc_lines = []
    while i < len(lines) and not is_section_header(lines[i]):
        line = lines[i].strip()
        if line and not line.startswith("목차") and not line.startswith("──"):
            toc_lines.append(line)
        i += 1

    if toc_lines:
        doc.add_heading("목차", level=1)
        for tl in toc_lines:
            p = doc.add_paragraph(tl)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            for run in p.runs:
                run.font.size = Pt(9)
        doc.add_page_break()

    # 본문 파싱
    current_section_title = None
    in_code = False
    code_lines = []
    table_lines = []
    in_table = False
    # 다이어그램 모드
    in_diagram = False
    diagram_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # === 구분선 건너뛰기
        if is_section_header(line):
            # 테이블/코드 블록 마무리
            if in_table and table_lines:
                make_table_from_lines(doc, table_lines)
                table_lines = []
                in_table = False
            if in_code and code_lines:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            if in_diagram and diagram_lines:
                add_code_block(doc, diagram_lines)
                diagram_lines = []
                in_diagram = False
            i += 1
            continue

        # 섹션 제목 (숫자. 또는 숫자-숫자. 로 시작)
        section_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        subsection_match = re.match(r"^(\d+)-(\d+)\.\s+(.+)$", stripped)
        subsubsection_match = re.match(r"^(\d+)-(\d+)-(\d+)\.\s+(.+)$", stripped)

        if subsubsection_match and not in_code and not in_table:
            title = f"{subsubsection_match.group(1)}-{subsubsection_match.group(2)}-{subsubsection_match.group(3)}. {subsubsection_match.group(4)}"
            doc.add_heading(title, level=3)
            i += 1
            continue

        if subsection_match and not in_code and not in_table:
            title = f"{subsection_match.group(1)}-{subsection_match.group(2)}. {subsection_match.group(3)}"
            doc.add_heading(title, level=2)
            i += 1
            continue

        if section_match and not in_code and not in_table:
            doc.add_heading(stripped, level=1)
            i += 1
            continue

        # ──── 소제목 구분선
        if is_subsection_divider(line) and not in_code:
            if in_table and table_lines:
                make_table_from_lines(doc, table_lines)
                table_lines = []
                in_table = False
            i += 1
            continue

        # 다이어그램 감지 (┌, │, └ 또는 화살표가 있는 ASCII art)
        if not in_table and not in_code:
            is_diagram_line = bool(
                re.search(r"[┌┐└┘│─▲▼▶◀►←→]", stripped)
                and not is_table_line(line)
            ) or (in_diagram and (stripped.startswith("|") or "──" in stripped or "│" in stripped or stripped == ""))

            # 테이블과 구분: 테이블은 │로 시작하고 데이터 셀이 있음
            if is_diagram_line and not re.match(r"^│\s*\S+\s*│\s*\S+\s*│", stripped):
                if not in_diagram:
                    in_diagram = True
                    diagram_lines = []
                if stripped:  # 빈 줄은 다이어그램 끝으로 간주하지만 중간 빈 줄은 유지
                    diagram_lines.append(line.rstrip())
                elif diagram_lines:
                    # 다이어그램 중간의 빈 줄
                    diagram_lines.append("")
                i += 1
                continue
            elif in_diagram:
                # 다이어그램 종료
                # 끝의 빈 줄 제거
                while diagram_lines and not diagram_lines[-1].strip():
                    diagram_lines.pop()
                if diagram_lines:
                    add_code_block(doc, diagram_lines)
                diagram_lines = []
                in_diagram = False

        # 테이블 처리
        if is_table_line(line) and not in_code:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # 테이블 종료
            if table_lines:
                make_table_from_lines(doc, table_lines)
            table_lines = []
            in_table = False

        # 코드 블록 감지 (들여쓰기 기반)
        indent = len(line) - len(line.lstrip()) if line.strip() else 0
        looks_like_code = (
            indent >= 2
            and stripped
            and (
                re.match(r"^(SELECT|MATCH|CREATE|MERGE|SET|RETURN|LOAD|WITH|WHERE|DELETE|REMOVE|OPTIONAL|UNWIND|ORDER|LIMIT)", stripped, re.IGNORECASE)
                or re.match(r"^(--|#|//|FROM|JOIN|AS\s|ON\s|AND\s|OR\s)", stripped, re.IGNORECASE)
                or re.match(r"^(def |class |import |from |for |if |elif |else:|try:|except|raise|return )", stripped)
                or re.match(r"^(docker |pip |git |RUN |COPY |ENV )", stripped)
                or re.match(r"^(CREATE |ALTER |DROP |INSERT |UPDATE )", stripped, re.IGNORECASE)
                or re.match(r"^\$\$", stripped)
                or re.match(r"^(escaped|full_query|result|cypher|s\s*=)", stripped)
                or re.match(r"^(\(|\)|\{|\})", stripped)
                or "execute_cypher" in stripped
                or "parse_agtype" in stripped
                or "cypher(" in stripped
                or ".replace(" in stripped
                or "agtype_build_map" in stripped
                or "shared_preload" in stripped
                or "search_path" in stripped
                or "json.loads" in stripped
                or "re.sub" in stripped
                or "isinstance" in stripped
                or stripped.startswith("for key")
                or stripped.startswith("holdings_data")
                or stripped.startswith("changes =")
                or stripped.startswith("auth_key")
                or stripped.startswith("client =")
                or stripped.startswith("stock.")
                or stripped.startswith("df =")
                or stripped.startswith("company =")
                or stripped.startswith("today_results")
                or stripped.startswith("yesterday_results")
            )
        )

        if looks_like_code:
            if not in_code:
                in_code = True
                code_lines = []
            code_lines.append(line.rstrip())
            i += 1
            continue
        elif in_code:
            # 빈 줄이면 아직 코드 블록 내부일 수 있음
            if not stripped:
                code_lines.append("")
                i += 1
                continue
            else:
                # 코드 블록 종료
                # 끝의 빈 줄 제거
                while code_lines and not code_lines[-1].strip():
                    code_lines.pop()
                if code_lines:
                    add_code_block(doc, code_lines)
                code_lines = []
                in_code = False

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 특수 포맷 처리

        # 주의/참고 블록
        if stripped.startswith("※"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(stripped)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(150, 80, 0)
            i += 1
            continue

        # 볼드 키-값 패턴: **key**: value 또는 (숫자) 번호 매기기
        bold_match = re.match(r"^-\s+\*\*(.+?)\*\*\s*[:\-]\s*(.+)$", stripped)
        if bold_match:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bold_match.group(1))
            run.font.bold = True
            p.add_run(f": {bold_match.group(2)}")
            i += 1
            continue

        # 번호 항목: (1), (2) 등
        num_match = re.match(r"^\((\d+)\)\s+(.+)$", stripped)
        if num_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"({num_match.group(1)}) ")
            run.font.bold = True
            p.add_run(num_match.group(2))
            i += 1
            continue

        # 불릿 항목: - 로 시작
        if stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        # 레이블 텍스트: "문제:", "해결:", "예시:" 등
        label_match = re.match(r"^(문제|해결|예시|상태|주의|보안 참고|이유|패턴|핵심 사항|핵심 가치|워크플로우|태그 예시|데이터 소스)[:\s]\s*$", stripped)
        if label_match:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.bold = True
            run.font.color.rgb = RGBColor(43, 87, 154)
            i += 1
            continue

        # 인라인 레이블: "문제: 내용..."
        inline_label = re.match(r"^(문제|해결|예시|상태|주의|보안 참고|이유|결과|변환 전|변환 후|잘못된 예|올바른 예|파싱 결과|오류 발생|감지):\s*(.+)$", stripped)
        if inline_label:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(f"{inline_label.group(1)}: ")
            run.font.bold = True
            run.font.color.rgb = RGBColor(43, 87, 154)
            p.add_run(inline_label.group(2))
            i += 1
            continue

        # 일반 텍스트
        p = doc.add_paragraph(stripped)
        i += 1

    # 잔여 블록 마무리
    if in_code and code_lines:
        while code_lines and not code_lines[-1].strip():
            code_lines.pop()
        if code_lines:
            add_code_block(doc, code_lines)
    if in_table and table_lines:
        make_table_from_lines(doc, table_lines)
    if in_diagram and diagram_lines:
        while diagram_lines and not diagram_lines[-1].strip():
            diagram_lines.pop()
        if diagram_lines:
            add_code_block(doc, diagram_lines)

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.save(str(OUTPUT))
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    build_docx()
